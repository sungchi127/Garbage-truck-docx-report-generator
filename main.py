import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os, sys
from paddleocr import PaddleOCR # Import PaddleOCR
from PIL import Image, ImageTk  # Import ImageTk
from docx import Document
from docx.shared import Inches
import re
import io  # Import io
import pandas as pd # Add pandas import

# --- Configuration ---
# You might need to set this if tesseract is not in your PATH
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe' # Example path
TEMPLATE_DIR = "templates"
OUTPUT_DIR = "output"
YELLOW_TEMPLATE = os.path.join(TEMPLATE_DIR, "template_yellow.docx")
WHITE_TEMPLATE = os.path.join(TEMPLATE_DIR, "template_white.docx")
MAPPING_FILE = os.path.join("license_mapping", "車牌對照表、輪胎規格表114.03.03.xlsx") # Path to mapping file

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(TEMPLATE_DIR, exist_ok=True) # Ensure template dir exists too



# --- Initialize PaddleOCR ---
# This needs to run only once to download and load model into memory
# Use lang='ch' for Chinese and English. use_gpu=False avoids needing CUDA setup.
print("Initializing PaddleOCR... This might take a moment on first run.")
ocr_engine = PaddleOCR(
    use_angle_cls=True,
    lang='ch',
    use_gpu=False,
    det_model_dir=r"C:\paddle_models\det\ch\ch_PP-OCRv4_det_infer",
    rec_model_dir=r"C:\paddle_models\rec\ch\ch_PP-OCRv4_rec_infer",
    cls_model_dir=r"C:\paddle_models\cls\ch_ppocr_mobile_v2.0_cls_infer"
)
print("PaddleOCR Initialized.")

# --- Load License Plate Mapping ---
def load_license_mapping(filepath):
    """Loads license plate to code mapping from the specified Excel file."""
    mapping = {}
    try:
        # Try reading the specified sheet name first, then common fallbacks
        sheet_name_to_try = '1.2級車牌複製用'
        try:
            df = pd.read_excel(filepath, sheet_name=sheet_name_to_try, header=None) # Assuming no header
        except ValueError: # If the specific sheet is not found
            print(f"Warning: Sheet '{sheet_name_to_try}' not found. Trying '工作表1' or 'Sheet1'...")
            try:
                df = pd.read_excel(filepath, sheet_name='工作表1', header=None)
            except ValueError:
                try:
                    df = pd.read_excel(filepath, sheet_name='Sheet1', header=None)
                except ValueError:
                    print(f"Error: Could not find sheets '{sheet_name_to_try}', '工作表1', or 'Sheet1' in {filepath}")
                    return {}

        # --- Final Revised Mapping Logic --- 
        # Reads PLATE(CODE) format directly from columns B, D, F (indices 1, 3, 5)
        mapping = {} # Reset mapping here before filling
        relevant_columns = [1, 3, 5] # Columns B, D, F

        for col_idx in relevant_columns:
            if col_idx < df.shape[1]: # Check column exists
                for item in df[col_idx].dropna(): # Iterate through non-empty cells in the column
                    item_str = str(item).strip()
                    # Regex to extract PLATE(CODE) format - Handles both half/full width parentheses
                    match = re.match(r'([A-Z0-9-]+)\s*[(（](\d+)[)）]', item_str, re.IGNORECASE)
                    if match:
                        plate = match.group(1).upper()
                        code = match.group(2) # Digits only
                        mapping[plate] = code
                        # Add version without hyphen too
                        mapping[plate.replace('-', '')] = code
                    else:
                        # Keep this uncommented for debugging
                        print(f"Debug: Item '{item_str}' in column {col_idx+1} did not match PLATE(CODE) format.")
                        pass
        # ------------------------------------

        if not mapping:
            print("Warning: No license plate mappings were loaded. Check file path, sheet name, and column format.")
        else:
            print(f"Successfully loaded {len(mapping)} license plate mappings.")
        return mapping

    except FileNotFoundError:
        print(f"*********************************************************************")
        print(f"*** Error: Mapping file not found at '{os.path.abspath(filepath)}' ***")
        print(f"*** Please ensure the file exists and the path is correct. ***")
        print(f"*********************************************************************")
        return {} # Return empty mapping on file not found
    except Exception as e: # General except block for any other errors during loading/processing
        print(f"*****************************************************")
        print(f"*** Error loading license plate mapping: {e} ***")
        print(f"*** Check file integrity, sheet names, and format. ***")
        print(f"*****************************************************")
        return {} # Return empty mapping on other errors

license_plate_map = load_license_mapping(MAPPING_FILE)
# --------------------------------

# --- OCR Function using PaddleOCR ---
def extract_data_from_image(image_path):
    # Placeholder function for OCR extraction.
    try:
        # No specific image preprocessing needed initially for PaddleOCR, it handles much internally
        # img = Image.open(image_path)
        # img = img.convert('L')

        print(f"--- Running PaddleOCR on {os.path.basename(image_path)} ---")
        result = ocr_engine.ocr(image_path, cls=True)
        print("--- PaddleOCR Raw Result ---")
        # print(result) # Can be very verbose
        print("-----------------------------")

        # Process PaddleOCR result into a single text block
        text_lines = []
        if result and result[0]: # Check if result is valid and contains data
            for line_info in result[0]:
                # line_info is like [[[coords]], ('text', confidence)]
                text_lines.append(line_info[1][0])
        text = "\n".join(text_lines)

        print("--- Reconstructed Text ---")
        print(text)
        print("---------------------------")


        # Basic extraction attempts (only for Address and Date now)
        # Date (look for 年月日 pattern)
        date_match = re.search(r'(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', text)
        date_str = date_match.group(0) if date_match else ""

        # Address (very basic - look for 路 or 街 or 號) - Needs improvement
        # Corrected Regex: Matches Chinese characters + street type + numbers + 号/號
        address_match = re.search(r'([\u4e00-\u9fff]+(?:路|街|巷|弄)\s*[\d-]+(?:號|号))', text)
        address_str = address_match.group(1).strip() if address_match else ""

        # --- Add License Plate Extraction ---
        # Regex to find common Taiwanese license plates (e.g., ABC-1234, KEL-0283, 1234-AB, TXY-1234)
        # This regex tries to cover various formats, including older ones and newer ones.
        # It looks for patterns like AAA-NNNN, AAA-NNN.N, AA-NNNN, NNNN-AA, K(E/A/...)L-NNNN etc.
        # Making it less strict about exact character counts where applicable.
        plate_match = re.search(r'([A-Z]{2,3}[- ]?[0-9]{3,4})|([0-9]{3,4}[- ]?[A-Z]{2,3})', text, re.IGNORECASE)
        plate_str = plate_match.group(0).upper().replace(' ', '-') if plate_match else ""
        code_str = None # Initialize code as None
        if plate_str and license_plate_map: # Check if map is loaded
            # Look up in map (try with and without hyphen)
            code_str = license_plate_map.get(plate_str) or license_plate_map.get(plate_str.replace('-', ''))
        # ------------------------------------

        # REMOVED OCR FOR PLATE AND TYPE - These will be manual
        # truck_type = ...

        return {"address": address_str, "date": date_str, "plate": plate_str, "code": code_str}

    except FileNotFoundError:
        return {"error": "圖片檔案未找到"}
    except Exception as e:
        print(f"OCR Error: {e}")
        return {"error": f"OCR 處理失敗: {e}"}


def generate_word_doc(data, img_path1, img_path2, output_filename):
    # Placeholder function for generating the Word document.
    # Determine template based on color (assuming manual selection for now)
    # For now, let's default to yellow, will add color selection later
    template_path = YELLOW_TEMPLATE # Defaulting, need GUI element for this
    # TODO: Add logic to select template based on GUI input for color

    try:
        if not os.path.exists(template_path):
             messagebox.showerror("錯誤", f"模板檔案未找到: {template_path}")
             return

        document = Document(template_path)

        # --- Replace Text Placeholders --- ## YOU ARE HERE
        replacements = {
            "{{ADDRESS}}": data.get("address", "N/A"),
            "{{DATE}}": data.get("date", "N/A"),
            "{{LICENSE_PLATE}}": data.get("plate", "N/A"),
            # Checkboxes
            "{{CHECKBOX_COMPRESSION}}": "■" if data.get("type") == "壓縮式垃圾車" else "□",
            "{{CHECKBOX_RECYCLING}}": "■" if data.get("type") == "資源回收車" else "□",
        }

        # Replace in paragraphs
        for p in document.paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    inline = p.runs
                    # Replace runs containing key
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, value)
                            inline[i].text = text

        # Replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, value)
                                        inline[i].text = text


        # --- Replace Image Placeholders ---
        # We need to find the paragraph containing the placeholder, clear it, then add picture
        def replace_image_placeholder(doc, placeholder, img_path, width_inches=3.0):
            replaced = False
            for p in doc.paragraphs:
                if placeholder in p.text:
                    # Clear the placeholder text first
                    for run in p.runs:
                        run.text = run.text.replace(placeholder, '')
                    # Add the picture in the now empty paragraph
                    try:
                        p.add_run().add_picture(img_path, width=Inches(width_inches))
                        # Optional: Center align the paragraph
                        # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        replaced = True # Placeholder found and replaced
                    except FileNotFoundError:
                         messagebox.showerror("錯誤", f"圖片檔案未找到: {img_path}")
                         return False
                    except Exception as e:
                         messagebox.showerror("錯誤", f"插入圖片時發生錯誤: {e}")
                         return False
            # Check tables too, if placeholders might be there
            for table in doc.tables:
                 for row in table.rows:
                     try: # Outer try for row.cells access
                         for cell in row.cells:
                             for p in cell.paragraphs:
                                 if placeholder in p.text:
                                     # Clear placeholder and attempt to insert image
                                     p.text = p.text.replace(placeholder, '') # Basic clear
                                     try: # Inner try for image insertion
                                         run = p.add_run()
                                         run.add_picture(img_path, width=Inches(width_inches))
                                         replaced = True # Set flag to True on success
                                     except FileNotFoundError:
                                         messagebox.showerror("錯誤", f"圖片檔案未找到: {img_path}")
                                     except Exception as e:
                                          messagebox.showerror("錯誤", f"插入表格圖片 '{os.path.basename(img_path)}' 時發生錯誤: {e}")
                                     # Break inner loops if image replaced
                                     if replaced: break # break paragraph loop
                             if replaced: break # break cell loop
                     except AttributeError:
                         # Skip rows that cause AttributeError (likely malformed)
                         print(f"Warning: Skipping a row in table due to unexpected structure (no 'cells' attribute).")
                         continue # Continue to the next row

                     # Break outer loop (row loop) if image replaced
                     if replaced: break
                 # Break table loop if image replaced
                 if replaced: break

            # Return status after checking both paragraphs and tables
            if not replaced:
                 print(f"Warning: Image placeholder '{placeholder}' not found in the document.")
            return replaced

        if img_path1:
            replace_image_placeholder(document, "{{IMAGE_1}}", img_path1, width_inches=5.0) # Adjust width as needed
        if img_path2:
            replace_image_placeholder(document, "{{IMAGE_2}}", img_path2, width_inches=5.0) # Adjust width as needed


        # Save the document
        doc_path = os.path.join(OUTPUT_DIR, output_filename)
        document.save(doc_path)
        messagebox.showinfo("成功", f"報告已產生: {doc_path}")

    except Exception as e:
        messagebox.showerror("錯誤", f"產生 Word 文件時發生錯誤: {e}")


# --- GUI Application ---
class App:
    def __init__(self, root):
        self.root = root
        self.root.title("垃圾車記錄產生器 (PaddleOCR)") # Updated title

        self.img_path1 = tk.StringVar()
        self.img_path2 = tk.StringVar()
        self.ocr_data = {} # To store combined results from OCR (mainly for potential pre-fill)

        # --- Manual Input Variables ---
        self.plate_var = tk.StringVar()
        self.address_var = tk.StringVar()
        self.date_var = tk.StringVar()
        self.truck_type_var = tk.StringVar(value="壓縮式垃圾車") # Restore truck type selection

        # Frame for Image 1 selection and preview
        frame1 = tk.LabelFrame(root, text="照片一 (通常包含地址/日期)", padx=5, pady=5)
        frame1.pack(padx=10, pady=5, fill="x")

        tk.Button(frame1, text="選擇照片一", command=lambda: self.select_image(self.img_path1, self.img_preview1_label)).pack(side=tk.LEFT, padx=5)
        tk.Label(frame1, textvariable=self.img_path1).pack(side=tk.LEFT, fill="x", expand=True)
        self.img_preview1_label = tk.Label(frame1) # Label to hold the image preview
        self.img_preview1_label.pack(pady=5)


        # Frame for Image 2 selection and preview
        frame2 = tk.LabelFrame(root, text="照片二 (通常包含車牌/車種)", padx=5, pady=5)
        frame2.pack(padx=10, pady=5, fill="x")

        tk.Button(frame2, text="選擇照片二", command=lambda: self.select_image(self.img_path2, self.img_preview2_label)).pack(side=tk.LEFT, padx=5)
        tk.Label(frame2, textvariable=self.img_path2).pack(side=tk.LEFT, fill="x", expand=True)
        self.img_preview2_label = tk.Label(frame2) # Label to hold the image preview
        self.img_preview2_label.pack(pady=5)

        # --- Frame for Manual Input Fields ---
        frame_manual_input = tk.LabelFrame(root, text="手動輸入/確認資訊", padx=10, pady=10)
        frame_manual_input.pack(padx=10, pady=10, fill="x")

        # Grid layout for better alignment inside this frame
        frame_manual_input.columnconfigure(1, weight=1) # Make entry column expandable

        # License Plate
        tk.Label(frame_manual_input, text="車牌號碼:").grid(row=0, column=0, padx=5, pady=2, sticky="w")
        plate_entry = tk.Entry(frame_manual_input, textvariable=self.plate_var, width=40)
        plate_entry.grid(row=0, column=1, padx=5, pady=2, sticky="ew")

        # Address
        tk.Label(frame_manual_input, text="檢查地點路段:").grid(row=1, column=0, padx=5, pady=2, sticky="w")
        address_entry = tk.Entry(frame_manual_input, textvariable=self.address_var, width=40)
        address_entry.grid(row=1, column=1, padx=5, pady=2, sticky="ew")

        # Date
        tk.Label(frame_manual_input, text="檢查日期:").grid(row=2, column=0, padx=5, pady=2, sticky="w")
        date_entry = tk.Entry(frame_manual_input, textvariable=self.date_var, width=40)
        date_entry.grid(row=2, column=1, padx=5, pady=2, sticky="ew")

        # Truck Type - Restore this section for template selection
        tk.Label(frame_manual_input, text="車種 (選擇模板):").grid(row=3, column=0, padx=5, pady=5, sticky="w") # Label changed slightly
        truck_type_frame = tk.Frame(frame_manual_input) # Inner frame for radio buttons
        truck_type_frame.grid(row=3, column=1, padx=5, pady=2, sticky="w")
        tk.Radiobutton(truck_type_frame, text="壓縮式垃圾車 (黃)", variable=self.truck_type_var, value="壓縮式垃圾車").pack(side=tk.LEFT) # Added color hint
        tk.Radiobutton(truck_type_frame, text="資源回收車 (白)", variable=self.truck_type_var, value="資源回收車").pack(side=tk.LEFT, padx=10) # Added color hint

        # Frame for Displaying OCR results (optional, could be integrated better)
        frame_results = tk.LabelFrame(root, text="OCR 辨識結果 (參考用)", padx=5, pady=5)
        frame_results.pack(padx=10, pady=5, fill="x")

        # Add Scrollbar to the results Text widget
        scrollbar = tk.Scrollbar(frame_results)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.result_text = tk.Text(frame_results, height=6, width=50, state=tk.DISABLED, yscrollcommand=scrollbar.set)
        self.result_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True) # Use fill/expand

        scrollbar.config(command=self.result_text.yview) # Link scrollbar to text widget

        # Generate Button
        tk.Button(root, text="產生報告", command=self.generate_report, font=('Arial', 12, 'bold')).pack(pady=20)

    def select_image(self, path_var, preview_label):
        file_path = filedialog.askopenfilename(
            title="選擇圖片",
            filetypes=[("Image Files", "*.png *.jpg *.jpeg *.bmp *.tiff")]
        )
        if file_path:
            path_var.set(file_path)
            self.display_image_preview(file_path, preview_label)
            # Immediately try OCR when an image is selected
            self.run_ocr_on_selection()

    def display_image_preview(self, file_path, preview_label, max_width=200, max_height=150):
        # Displays a preview of the selected image in the GUI.
        try:
            img = Image.open(file_path)
            img.thumbnail((max_width, max_height)) # Resize preserving aspect ratio
            photo = ImageTk.PhotoImage(img)

            preview_label.config(image=photo)
            preview_label.image = photo # Keep a reference! Important for Tkinter.
        except Exception as e:
            preview_label.config(text=f"無法預覽:\n{e}", image='')
            preview_label.image = None # Clear reference


    def run_ocr_on_selection(self):
        # Runs OCR on selected images and updates the results display.
        self.ocr_data = {} # Reset previous results
        results_display = ""
        ocr_address = ""
        ocr_date = ""
        ocr_plate = "" # Add variable for plate
        ocr_code = ""  # Add variable for code

        path1 = self.img_path1.get()
        path2 = self.img_path2.get()

        if path1:
            data1 = extract_data_from_image(path1)
            if "error" not in data1:
                 # Store potential OCR results
                 if data1.get('address'): ocr_address = data1.get('address')
                 if data1.get('date'): ocr_date = data1.get('date')
                 if data1.get('plate'): ocr_plate = data1.get('plate') # Get plate if found
                 if data1.get('code'): ocr_code = data1.get('code')    # Get code if found
                 plate_display = f"{ocr_plate}({ocr_code})" if ocr_plate and ocr_code else ocr_plate # Format for display
                 results_display += f"照片一 OCR (參考):\n  地址: {ocr_address or 'N/A'}\n  日期: {ocr_date or 'N/A'}\n  車牌: {plate_display or 'N/A'}\n" # Add formatted plate to display
            else:
                 results_display += f"照片一 OCR 錯誤: {data1['error']}\n"


        if path2:
            data2 = extract_data_from_image(path2)
            if "error" not in data2:
                 # Update address/date only if first image didn't find it
                 if not ocr_address and data2.get('address'): ocr_address = data2.get('address')
                 if not ocr_date and data2.get('date'): ocr_date = data2.get('date')
                 # Update plate/code only if not found in first image
                 if not ocr_plate and data2.get('plate'):
                     ocr_plate = data2.get('plate')
                     ocr_code = data2.get('code') # Also update code if plate is updated
                 elif ocr_plate and not ocr_code and data2.get('code'): # If plate found in img1 but code not, try img2 for code
                     ocr_code = data2.get('code')

                 plate_display = f"{ocr_plate}({ocr_code})" if ocr_plate and ocr_code else ocr_plate # Format for display
                 results_display += f"照片二 OCR (參考):\n  地址: {data2.get('address', 'N/A')}\n  日期: {data2.get('date', 'N/A')}\n  車牌: {plate_display or 'N/A'}\n" # Add formatted plate to display
            else:
                 results_display += f"照片二 OCR 錯誤: {data2['error']}\n"


        # --- Pre-fill manual fields with OCR results (if found) ---
        if ocr_address:
            self.address_var.set(ocr_address)
        # else: # Optional: Clear if OCR fails? Or leave previous manual input?
            # self.address_var.set("")
        if ocr_date:
            self.date_var.set(ocr_date)
        if ocr_plate: # Pre-fill plate if found
            plate_display = f"{ocr_plate}({ocr_code})" if ocr_code else ocr_plate # Format plate with code if available
            self.plate_var.set(plate_display)
        # else:
            # self.date_var.set("")
        # else:
            # self.plate_var.set("") # Optional: Clear plate if not found

        # Update the text display widget
        print("--- Updating GUI Text Area ---") # Debug print added
        print(f"Content to display:\\n{results_display}") # Debug print added
        print("-----------------------------") # Debug print added
        self.result_text.config(state=tk.NORMAL) # Enable editing
        self.result_text.delete('1.0', tk.END) # Clear previous text
        self.result_text.insert(tk.END, results_display)
        self.result_text.config(state=tk.DISABLED) # Disable editing


    def generate_report(self):
        img1 = self.img_path1.get()
        img2 = self.img_path2.get()

        if not img1 or not img2:
            messagebox.showwarning("警告", "請選擇兩張照片")
            return

        # --- Get data from MANUAL INPUT fields --- VITAL CHANGE
        manual_plate_full = self.plate_var.get().strip() # Get potentially formatted plate
        manual_address = self.address_var.get().strip()
        manual_date = self.date_var.get().strip()
        manual_truck_type = self.truck_type_var.get() # Restore getting truck type

        # --- Extract only the plate number from the potentially formatted input ---
        plate_parse_match = re.match(r'([A-Z0-9-]+)(?:s*(\(\d+))?$', manual_plate_full, re.IGNORECASE)
        manual_plate = plate_parse_match.group(1).upper() if plate_parse_match else manual_plate_full # Fallback to full string if parse fails
        # -----------------------------------------------------------------------

        # Basic validation (optional but recommended)
        if not manual_plate:
            messagebox.showwarning("警告", "請輸入車牌號碼")
            return
        if not manual_address:
            messagebox.showwarning("警告", "請輸入檢查地點路段")
            return
        if not manual_date:
            messagebox.showwarning("警告", "請輸入檢查日期")
            return
        # Truck type validation not strictly needed as it has a default

        final_data = {
            "plate": manual_plate,
            "address": manual_address,
            "date": manual_date,
            # Type is used for template selection, not replacement
        }
        # -------------------------------------------

        # --- Extract the 3-digit code directly from the full input string ---
        code_match = re.search(r'[(（](\d+)[)）]', manual_plate_full)
        if code_match:
            plate_code_3digit = code_match.group(1)
            print(f"Extracted code '{plate_code_3digit}' directly from input '{manual_plate_full}'.")
        else:
            plate_code_3digit = "XXX" # Use placeholder if no code in parentheses found in input
            print(f"Warning: No code found in parentheses in input '{manual_plate_full}'. Using 'XXX' in filename.")
        # ---------------------------------------------------------------------

        # Create a meaningful output filename using the new format
        # --- Generate filename based on truck type and extracted code ---
        if manual_truck_type == "壓縮式垃圾車": # Corresponds to Yellow Template
            output_filename = f"空白-1.2級檢查-{plate_code_3digit}垃圾車.docx"
            template_path = YELLOW_TEMPLATE
            print(f"選擇模板: {YELLOW_TEMPLATE}")
        elif manual_truck_type == "資源回收車": # Corresponds to White Template
            output_filename = f"空白-1.2級檢查-{plate_code_3digit}回收車.docx"
            template_path = WHITE_TEMPLATE
            print(f"選擇模板: {WHITE_TEMPLATE}")
        else:
            # Handle unexpected case or default - Use a fallback name
            messagebox.showerror("錯誤", f"未知的車種選擇: {manual_truck_type}. 無法產生檔名。")
            # Fallback filename if type is unknown
            # output_filename = f"{manual_plate.replace('-', '')}_XXX_report.docx" # Example fallback
            # template_path = YELLOW_TEMPLATE # Default template maybe? Or just return
            print(f"錯誤: 未知的車種選擇: {manual_truck_type}. 無法決定模板與檔名。")
            return # Stop processing if truck type is invalid

        print(f"輸出檔名將為: {output_filename}")
        # -------------------------------------------

        # --- Generate the Word Document ---
        try:
            if not os.path.exists(template_path):
                 messagebox.showerror("錯誤", f"模板檔案未找到: {template_path}")
                 return

            document = Document(template_path)

            # --- Replace Text Placeholders using MANUAL data --- ## YOU WERE HERE
            replacements = {
                "{{ADDRESS}}": final_data.get("address", "N/A"),
                "{{DATE}}": final_data.get("date", "N/A"),
                "{{LICENSE_PLATE}}": final_data.get("plate", "N/A"),
                # Checkboxes are now static in the template, no replacement needed
            }

            # Replace in paragraphs (Restore run-level replacement logic)
            for p in document.paragraphs:
                for key, value in replacements.items():
                    if key in p.text:
                        inline = p.runs
                        # Replace runs containing key while preserving style
                        for i in range(len(inline)):
                            if key in inline[i].text:
                                text = inline[i].text.replace(key, value)
                                inline[i].text = text


            # Replace in tables (Restore and enable run-level replacement logic)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in p.text:
                                    inline = p.runs
                                    # Replace runs containing key while preserving style
                                    for i in range(len(inline)):
                                        if key in inline[i].text:
                                            text = inline[i].text.replace(key, value)
                                            inline[i].text = text
            # ----------------------------------------------------------


            # --- Replace Image Placeholders ---
            def replace_image_placeholder(doc, placeholder, img_path, width_inches=3.0):
                replaced = False
                for p in doc.paragraphs:
                    if placeholder in p.text:
                        # Clear placeholder and insert image
                        p.text = p.text.replace(placeholder, '') # Basic clear
                        try:
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(width_inches))
                            replaced = True
                        except FileNotFoundError:
                            messagebox.showerror("錯誤", f"圖片檔案未找到: {img_path}")
                        except Exception as e:
                            messagebox.showerror("錯誤", f"插入圖片 '{os.path.basename(img_path)}' 時發生錯誤: {e}")
                        # Break after first replacement in paragraphs if needed, or handle multiple occurrences
                        # For now, assume unique placeholders
                        if replaced: break
                if replaced: return True

                # Check tables if not found in paragraphs
                for table in doc.tables:
                     for row in table.rows:
                         try: # Outer try for row.cells access
                             for cell in row.cells:
                                for p in cell.paragraphs:
                                     if placeholder in p.text:
                                         # Clear placeholder and attempt to insert image
                                         p.text = p.text.replace(placeholder, '') # Basic clear
                                         try: # Inner try for image insertion
                                             run = p.add_run()
                                             run.add_picture(img_path, width=Inches(width_inches))
                                             replaced = True # Set flag to True on success
                                         except FileNotFoundError:
                                             messagebox.showerror("錯誤", f"圖片檔案未找到: {img_path}")
                                         except Exception as e:
                                              messagebox.showerror("錯誤", f"插入表格圖片 '{os.path.basename(img_path)}' 時發生錯誤: {e}")
                                         # Break inner loops if image replaced
                                         if replaced: break # break paragraph loop
                             if replaced: break # break cell loop
                         except AttributeError:
                             # Skip rows that cause AttributeError (likely malformed)
                             print(f"Warning: Skipping a row in table due to unexpected structure (no 'cells' attribute).")
                             continue # Continue to the next row

                         # Break outer loop (row loop) if image replaced
                         if replaced: break
                     # Break table loop if image replaced
                     if replaced: break

                # Return status after checking both paragraphs and tables
                if not replaced:
                     print(f"Warning: Image placeholder '{placeholder}' not found in the document.")
                return replaced


            if img1 and not replace_image_placeholder(document, "{{IMAGE_1}}", img1, width_inches=5.0):
                 print(f"Warning: Image placeholder '{{IMAGE_1}}' not found in the document.")
            if img2 and not replace_image_placeholder(document, "{{IMAGE_2}}", img2, width_inches=5.0):
                 print(f"Warning: Image placeholder '{{IMAGE_2}}' not found in the document.")


            # Save the document
            doc_path = os.path.join(OUTPUT_DIR, output_filename)
            document.save(doc_path)
            messagebox.showinfo("成功", f"報告已產生於:\n{os.path.abspath(doc_path)}")

        except Exception as e:
            messagebox.showerror("錯誤", f"產生 Word 文件時發生嚴重錯誤: {e}")


# --- Main Execution ---
if __name__ == "__main__":
    # Check if template files exist
    if not os.path.exists(YELLOW_TEMPLATE):
        with open(YELLOW_TEMPLATE, 'w') as f: # Create dummy if not exists
            f.write("這是黃色卡車模板。請替換為您的 .docx 模板。\n")
            f.write("{{ADDRESS}}\n{{DATE}}\n{{LICENSE_PLATE}}\n")
            f.write("壓縮式垃圾車: {{CHECKBOX_COMPRESSION}}\n")
            f.write("資源回收車: {{CHECKBOX_RECYCLING}}\n")
            f.write("{{IMAGE_1}}\n{{IMAGE_2}}\n")
        print(f"Warning: Created dummy template file at {YELLOW_TEMPLATE}")

    if not os.path.exists(WHITE_TEMPLATE):
         with open(WHITE_TEMPLATE, 'w') as f: # Create dummy if not exists
            f.write("這是白色卡車模板。請替換為您的 .docx 模板。\n")
            f.write("{{ADDRESS}}\n{{DATE}}\n{{LICENSE_PLATE}}\n")
            f.write("壓縮式垃圾車: {{CHECKBOX_COMPRESSION}}\n")
            f.write("資源回收車: {{CHECKBOX_RECYCLING}}\n")
            f.write("{{IMAGE_1}}\n{{IMAGE_2}}\n")
         print(f"Warning: Created dummy template file at {WHITE_TEMPLATE}")


    root = tk.Tk()
    app = App(root)
    root.mainloop() 

# 打包指令
